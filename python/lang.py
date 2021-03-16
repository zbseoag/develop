#!/usr/bin/env python
# -*- coding: utf8 -*-
# if __name__ == '__main__':

PY_PYTHON=3
#默认库将驻留在 C:\Python\Lib\ 中，第三方模块存储在 C:\Python\Lib\site-packages\
PYTHONPATH

#数学计算
/   除运算返回浮点数
//  舍去法取整
%   取余数
**  双星号是乘方
表达式被赋值给变量 _

Decimal  小数
Fraction 分数
complex 复数

r'C:\some\name' 字符串前面加 "r" 原样输出

# 三引号可以包含多行字符块,也包含换行,如果不想包含换行,加 "\"
print("""\
Usage: thingy [OPTIONS]
     -h                        Display this usage message
     -H hostname               Hostname to connect to
""")

#字符串拼接用加号,字符串乘数表重复次数
a =  'a' + 'x' * 3

#自动拼接
b = 'abc' 'def'
('abc'
'def')

字符串中的单个字符可以通过索引访问,负值索引表倒数位置,由于 -0 和 0 是一样的,表示第一位,所以负数索引从 -1 开始
word[-1]


开始总是被包括在结果中，而结束不被包括。这使得 s[:i] + s[i:] 总是等于 s
从索引2到5,即:第三位到第五位, 索引5表示第六位
word[2:5]
倒数-1表示倒数第一,即负几表示倒数第几
word[-2:]

切片中的越界索引会被自动处理,不会报错
word[4:42]

字符串不能被修改,向字符串某个索引位置赋值

支持相加合并操作
squares + [36, 49, 64, 81, 100]
list.append(216)
list[2:5] = ['C', 'D', 'E']
letters[2:5] = [] #清空第三到第五个元素

注释符号：#  '''  """

模块导入语法
将整个模块导入： import somemodule
从模块中导入函数： from <module> import <function>
从模块中导入多个函数： from <module> import <function1>, <function2>
从模块中导入所有函数格： from <module> import *

查看关键字
import keyword
keyword.kwlist

多个变量赋值
a, b, c = 1, 2, "string"

数据类型
Number（数字）
String（字符串）
List（列表）：list=[]
Tuple（元组）：tup=(), tup=(20,)
Set（集合）：set=set()
Dictionary（字典）：dict={}

不可变数据类型：Number、String、Tuple
可变数据类型：List、Set、Dictionary

数据类型转换
int(x [,base])：将x转换为整数
float(x)：将x转换到一个浮点数
complex(real [,imag])：创建一个复数
str(x)：转换为字符串
repr(x)：转换为表达式字符串
eval(str)：执行有效Python表达式，如 eval( '3 * x' )
tuple(s)：转换为元组，如 tuple([1,2,3]), tuple({'a':1, 'b':2})
list(s)：转换为列表,如 list((1,2,3))
set(s)：转换集合,如 set([1,2,3])
dict(d)：创建一个字典,如 dict(a=1, b=2) dict(['a','b'], [1, 2]) dict([('a', 1), ('b', 2)])
frozenset(s)：转换为不可变集合
chr(x)：整数转换 ASCII 字符
ord(x)：转换为ASCII 数值
hex(x)：转十六进制
oct(x)：转八进制


运算符:
算术运算符：四则运算 % **幂  //整除（舍去法)
比较运算符：== !=  > <  >=  <=
赋值运算符：= += -=  *=  /= %=  **= //=  :=
逻辑运算符：and，or， not
位运算符：& | ^ ~ << >>
           成员运算符：in，not in
身份运算符：is，is not

转义字符
\ (在行尾时)续行符
\a	响铃
\b	退格
\000 空
\v	纵向制表符
\f	换页
\o	八进制数
\x	十六进制数


ascii(object)   返回一个对象可打印的字符串
bin(x)   转二进制, 若 x 不是 int 对象，那它需要定义 __index__() 方法来返回一个整数
class bool(x)   返回一个布尔值, 默认返回 False
breakpoint(*args, **kws)  此函数会在调用时将你陷入调试器中。
class bytearray([source[, encoding[, errors]]])  返回一个新的 bytes 数组。 bytearray 类是一个可变序列，包含范围为 0 <= x < 256 的整数。
class bytes([source[, encoding[, errors]]])  返回一个新的“bytes”对象， 是一个不可变序列，包含范围为 0 <= x < 256 的整数。
callable(object)   对象是否可调用的,如果实例的类有 __call__() 方法，则它是可调用。
chr(i) 返回 Unicode 字符
ord(c) 返回代表 Unicode 字符的码点整数
@classmethod 把一个方法封装成类方法。
compile(source, filename, mode, flags=0, dont_inherit=False, optimize=-1
class complex([real[, imag]])  返回值为 real + imag*1j 的复数，或将字符串或数字转换为复数。
delattr(object, name)  删除对象属性
class dict(**kwarg)
class dict(mapping, **kwarg)
class dict(iterable, **kwarg) 创建一个新的字典。
dir([object])  如果没有实参，则返回当前本地作用域中的名称列表。如果有实参，它会尝试返回该对象的有效属性列表。
divmod(a, b)  整除,并返回一对商和余数。
eval(expression[, globals[, locals]])实参是一个字符串，以及可选的 globals 和 locals。
filter(function(item),  iterable) 用回调函数 过滤 iterable 中元素
class frozenset([iterable])返回一个新的 frozenset 对象，它包含可选参数 iterable 中的元素。
getattr(object, name[, default])返回对象的属性值, 若属性不存在，则返回 default 值
globals()返回表示当前全局符号表的字典。
hasattr(object, name) 对象是否具有某个属性。
enumerate(iterable, start=0)返回一个枚举对象。iterable 必须是一个序列，或 iterator，或其他支持迭代的对象。
exec(object[, globals[, locals]])这个函数支持动态执行 Python 代码。
help([object])启动内置的帮助系统（此函数主要在交互式中使用）
hex(x)将整数转换为以“0x”为前缀的小写十六进制字符串
id(object) 返回对象的“标识值”。该值是一个整数，在此对象的生命周期中保证是唯一且恒定的。
input([prompt]) 如果存在 prompt 实参，则将其写入标准输出，末尾不带换行符。接下来，该函数从输入中读取一行，将其转换为字符串（除了末尾的换行符）并返回。
class int(x, base=10)
    isinstance(object, classinfo) object 对象就否是 classinfo 的实例
issubclass(class, classinfo)  class 是否是 classinfo 的子类
iter(object[, sentinel])  返回一个 iterator 对象。
len(s) 返回对象的长度或元素个数。
class list([iterable])  虽然被称为函数，list 实际上是一种可变序列类型，
locals() 更新并返回表示当前本地符号表的字典。
map(function(item1, item2), iterable1, iterable2...)   返回一个将 function 应用于 iterable 中每一项并输出其结果的迭代器
max(iterable, *[, key, default]) max(arg1, arg2, *args[, key]) 返回可迭代对象中最大的元素，或者返回两个及以上实参中最大的。
min(iterable, *[, key, default]) min(arg1, arg2, *args[, key]) 返回可迭代对象中最小的元素，或者返回两个及以上实参中最小的。
memoryview(obj) 返回由给定实参创建的“内存视图”对象。
next(iterator[, default]) 通过调用 iterator 的 __next__() 方法获取下一个元素。如果迭代器耗尽，则返回给定的 default
class object 返回一个没有特征的新对象。
oct(x)将一个整数转变为一个前缀为“0o”的八进制字符串。
open(file, mode='r', buffering=-1, encoding=None, errors=None, newline=None, closefd=True, opener=None)   打开 file 并返回对应的 file object。
pow(x, n[, z]) 返回 x 的 n 次幂；如果 z 存在，则对 z 取余
print(*objects, sep=' ', end='\n', file=sys.stdout, flush=False)  将 objects 打印到 file 指定的文本流
range(stop) range(start, stop[, step]) 虽然被称为函数，但 range 实际上是一个不可变的序列类型
repr(object) 返回包含一个对象的可打印表示形式的字符串。
reversed(seq) 返回一个反向的 iterator。 seq 必须是一个具有 __reversed__() 方法的对象或者是支持该序列协议（具有从 0 开始的整数类型参数的 __len__() 方法和 __getitem__() 方法）
round(number[, ndigits]) 返回 number 舍入到小数点后 ndigits 位精度的值
class set([iterable]) 返回一个新的 set 对象，可以选择带有从 iterable 获取的元素。 set 是一个内置类型
setattr(object, name, value)
class slice(stop) class slice(start, stop[, step]) 返回一个表示由 range(start, stop, step) 所指定索引集的 slice 对象
sorted(iterable, *, key=None, reverse=False) 根据 iterable 中的项返回一个新的已排序列表。
@staticmethod 将方法转换为静态方法
class str(object='') class str(object=b'', encoding='utf-8', errors='strict') 返回一个 str 版本的 object
sum(iterable[, start])
super([type[, object-or-type]])
class type(object) class type(name, bases, dict) 传入一个参数时，返回 object 的类型
vars([object]) 对象 __dict__ 属性
zip(*iterables) 创建一个聚合了来自每个可迭代对象中的元素的迭代器
__import__(name, globals=None, locals=None, fromlist=(), level=0)

删除变量：del var1, var2

条件语句
if(): elif(): else:
while(): else:
for in : else:

def func(var1, var2=Null, *varN, **kwargs):

def f(a,b,*,c, d):
    pass
f(a, b, c=3, d=4)


lambda 函数
sum = lambda arg1, arg2: arg1 + arg2

强制位置参数
def f(a, b, /, c, d, *, e, f): 形参 a 和 b 必须使用位置参数，c 或 d 可以是位置或关键字参数，而 e 或 f 必须为关键字参数:

sys.argv 是一个包含命令行参数的列表。
sys.path 包含了一个 Python 解释器自动查找所需模块的路径的列表。


打开文件的模式：读从文件开头起，重写会创建或清空，追加才在文件尾，附加属性更全能。如果不想清空原文件只有"r+"和"追加"能选，切记。
r	只读，指针定位文件头
r+	读附加写，指针定位文件头，写入内容会覆盖
w	重写，创建或清空文件，一切从新开始
w+	重写附加读，写完接着读
a	追加写入，可以创建文件
a+	追加写入，写完接着读

file = open(file, mode='r', buffering=-1, encoding=None, errors=None, newline=None, closefd=True, opener=None)
file.close() 关闭文件。
file.flush()刷新文件内部缓冲，直接把内部缓冲区的数据立刻写入文件, 而不是被动的等待输出缓冲区写入。
file.fileno()返回一个整型的文件描述符(file descriptor FD 整型), 可以用在如os模块的read方法等一些底层操作上。
file.isatty()如果文件连接到一个终端设备返回 True，否则返回 False。
file.read([size])从文件读取指定的字节数，如果未给定或为负则读取所有。
file.readline([size])读取整行，包括 "\n" 字符。
file.readlines([sizeint])读取所有行并返回列表，若给定sizeint>0，返回总和大约为sizeint字节的行, 实际读取值可能比 sizeint 较大, 因为需要填充缓冲区。
file.seek(offset[, whence])移动文件读取指针到指定位置
file.tell() 返回文件当前位置。
file.truncate([size])从文件的首行首字符开始截断，截断文件为 size 个字符，无 size 表示从当前位置截断；截断之后后面的所有字符被删除，其中 Widnows 系统下的换行代表2个字符大小。
file.write(str) 将字符串写入文件，返回的是写入的字符长度。
file.writelines(sequence) 向文件写入一个序列字符串列表，如果需要换行则要自己加入每行的换行符。

远程访问
from urllib.request import urlopen
for line in urlopen('http://tycho.usno.navy.mil/cgi-bin/timer.pl'):
    line = line.decode('utf-8')
    if 'EST' in line or 'EDT' in line:
        print(line)


import smtplib
server = smtplib.SMTP('localhost')
server.sendmail('soothsayer@example.org', 'jcaesar@example.org',
                """To: jcaesar@example.org
                From: soothsayer@example.org
                Beware the Ides of March.
                """)
server.quit()


性能度量
from timeit import Timer
Timer('t=a; a=b; b=t', 'a=1; b=2').timeit()
Timer('a,b = b,a', 'a=1; b=2').timeit()


单元测试
import unittest

class TestStatisticalFunctions(unittest.TestCase):

    def test_average(self):
        self.assertEqual(average([20, 30, 70]), 40.0)
        self.assertEqual(round(average([1, 5, 7]), 1), 4.3)
        self.assertRaises(ZeroDivisionError, average, [])
        self.assertRaises(TypeError, average, 20, 30, 70)

unittest.main()

import re
match = re.match(r'(.)是(.)', "你是猪", re.M|re.I)
match2 = re.search(r'(.)是(.)', "你是猪", re.M|re.I)
match.group()

#替换
num = re.sub(r'#.*', '', "2004-959-559 # 这是一个电话号码")

pattern = re.compile(r'\d+') # 用于匹配至少一个数字
m = pattern.match('one12twothree34four') # 查找头部，没有匹配

pattern = re.compile(r'\d+')   # 查找数字
result1 = pattern.findall('runoob 123 google 456')
result2 = pattern.findall('run88oob123google456', 0, 10)

返回可迭代对象
it = re.finditer(r"\d+","12a32bc43jf3")
for match in it:
    print (match.group() )

re.split('\W+', 'runoob, runoob, runoob.')



def test2():
    import time
    url = 'http://cdn.heweather.com/china-city-list.txt'
    html = requests.get(url)
    # 前6行就无用数据
    data = html.text.split('\n')[6:]

    document = tools.connect('weather', 'sheet_weather')
    for item in data[0:30]:
        # item = '| CN101010100' 前2位无用,再往后取11位,所以就就[2:13]
        url = 'https://free-api.heweather.com/v5/forecast?key=7d0daf2a85f64736a42261161cd3060b&city=' + item[2:13]
        html = requests.get(url)
        html.encoding = 'utf8'
        document.insert_one(html.json())

        tools.log(html.text)
        time.sleep(1)

def test3():
    import urllib
    import time
    url = 'https://touch.dujia.qunar.com/depCities.qunar'
    html = requests.get(url)
    text = html.json()
    for item in text['data']:
        for dep in text['data'][item]:
            a = []
            url = 'https://m.dujia.qunar.com/golfz/sight/arriveRecommend?dep={}&exclude=&extensionImg=255,175'.format(urllib.request.quote(dep))

            response = requests.get(url)
            arrive = response.json()
            for item in arrive['data']:
                for subitem in item['subModules']:
                    for query in subitem['items']:
                        if query['query'] not in a:
                            a.append(query['query'])
            print(a)
            exit()

def test4():
    import keyboard
    from PIL import ImageGrab
    import time

    if not keyboard.wait(hotkey='f1'):
        if not keyboard.wait(hotkey='ctrl+c'):
            time.sleep(0.01)
            image = ImageGrab.grabclipboard()
            image.save('test.png')

def find_record(city='北京'):
    document = tools.connect('weather', 'sheet_weather')
    for item in document.find({'HeWeather5.basic.city': city}):
        print(item)


def test5():
    import os
    os.environ['qikang'] = 'leslie'
    for k, v in os.environ.items():
        print(k, v)


import requests
import json
from bs4 import BeautifulSoup
import re
import math
import tools


def test6():

    l = [1, 2, 3];      p(l * 5)
    s = 'abc-';     p(5 * s)
    board = [['-'] * 3 for i in range(3)];    p(board)
    board[1][2] = 'X';      p(board)

    #等同如下代码
    board = []
    for i in range(3):
        row = ['-'] * 3
        board.append(row)
    p(board)
    board[1][2] = 'X';  p(board)


    weird = [['-'] * 3] * 3; p(weird) # 最外层列表，其实包含三个指向同一个列表的引用
    weird[1][2] = '0';  p(weird)    # 当只改变其中一个列表中的元素时，才发现都发生了改变

    # 等同如下代码
    row = ['-'] * 3
    board = []
    for i in range(3):
        board.append(row)
    board[1][2] = 'Y';  p(board)

def test7():
    l = [1, 2, 3]; p(id(l))
    l *= 2; p(id(l)) # 运用增量乘法后，列表的 ID 没变，新元素追加到列表上

    t = (1, 2, 3);  p(id(t))
    t *= 2;     p(id(t)) # 运用增量乘法，新的无组被创建

def test9():

    import pprint
    t = [[[['black', 'cyan'], 'white', ['green', 'red']], [['magenta', 'yellow'], 'blue']]]
    pprint.pprint(t, width=30)

    # textwrap 模块能够格式化文本段落，以适应给定的屏幕宽度:
    import textwrap
    doc = """The wrap() method is just like fill() except that it returns
    a list of strings instead of one big string with newlines to separate the wrapped lines."""

    print(textwrap.fill(doc, width=40))

    from string import Template
    t = Template('${village} folk send $$10 to $cause.')
    a = t.substitute(village='Nottingham', cause='the ditch fund'); print(a)

    t = Template('Return the $item to $owner.')
    t.substitute(dict(item='unladen swallow'))
    # 如果数据缺失，safe_substitute 会直接将占位符原样保留。
    t.safe_substitute({})

def test8():

    # 列表推导式
    s = list(map(lambda x: x ** 2, range(10)))
    s = [x ** 2 for x in range(10)]
    b = [(x, y, z) for x in [1, 2, 3] for y in [3, 1, 4] for z in [8, 9, 10] if x != y]

    matrix = [
        [1, 2, 3, 4],
        [5, 6, 7, 8],
        [9, 10, 11, 12],
    ]
    abc = list(zip(*matrix)); p(abc)

    a = set('1235'); b = set('124'); print(a, b, a - b);

def test9():
    import functools
    int2 = functools.partial(int, base=2)
    int2(10, 2)

    max2 = functools.partial(max, 10)
    max2(1, 2, 3)


def test10():

    def lazy_sum(*args):

        def sum():
            ax = 0
            for n in args:
                ax = ax + n
                return ax

        return sum

    f = lazy_sum(1, 2, 3, 4)
    a = f()
    print(a)


def test11():
    import functools

    # 生成器
    generator = (x * x for x in range(10))

    # map 用回调函数处理每一个元素
    a = map(str, [1, 2, 3]); p(a)
    a = functools.reduce(lambda x, y: x+y, [1,2,3]); p(a)

    def is_odd(n):
        return n % 2 == 1
    # 过滤列表元素
    a = filter(is_odd, [1, 2, 3, 4, 5, 6]); print(a)

    # 排序时，先求绝对值
    a = sorted([1, 2, -3, -4, -5, 6], key=abs); p(a)
    a = sorted(['bob', 'about', 'Zone', 'Credit'], key=str.lower, reverse=True); p(a)



def test12():

    def f(ham: str, eggs: str = 'eggs') -> str:
        """这是一个测试
            :param ham:
            :param eggs:
            :return:
        """
        return ham + ' ' + eggs

    f('spam')
    print( f.__doc__, f.__annotations__)


def test13():

    def fib2(n):
        """Return a list containing the Fibonacci series up to n."""
        result = []
        a, b = 0, 1
        while a < n:
            result.append(a)
            a, b = b, a + b
        return result

    f100 = fib2(100); p(f100)

    import logging
    logging.debug('Debugging information')

def test14():

    #数值精度
    import decimal

    a = decimal.getcontext();   p(a)
    a = decimal.Decimal(1) / decimal.Decimal(9);    p(a)
    decimal.getcontext().prec = 4 #精度改为保留4位小数
    a = decimal.Decimal(1) / decimal.Decimal(9);    p(a)

def test15():
    # 读取文件
    with open("data.txt") as f:
        for line in f:
            print(line, end="")

def test16():
    # 队列
    import collections
    queue = collections.deque(["a", "b", "c"])
    queue.append("d")
    queue.popleft();    p(queue)

def test17():
    # 操作文件
    import shutil
    shutil.copyfile('aaa.txt', 'bbb.txt')

    for line in open("myfile.txt"):
        print(line, end='')

    with open('aaa.txt') as f:
        # 读取所有内容
        content = f.read()
        content = f.readlines()


def test18():
    # 系统管理
    import os
    os.getcwd()  # 当前工作目录
    os.chdir('/server/accesslogs')  # Change current working directory
    os.system('mkdir today')
    os.listdir('.')

    import sys
    print(sys.path, sys.argv)
    sys.stderr.write('Warning, log file not found starting a new one\n')
    sys.path.append('/ufs/guido/lib/python')
    sys.exit()



def test19():
    # 文件名匹配库
    import glob
    glob.glob('*.py')


def test20():
    # 计算数值数据的基本统计属性（均值，中位数，方差等）
    import statistics
    data = [2.75, 1.75, 1.25, 0.25, 0.5, 1.25, 3.5]
    a = statistics.mean(data); p(a)
    a = statistics.median(data);p(a)
    a = statistics.variance(data);p(a)

def test21():

    #遍历二维列表
    list = [[1, 2, 3], [4, 5, 6], [7, 8, 9]]
    a = [n for row in list for n in row];
    print(a)


def test22():
    import smtplib
    server = smtplib.SMTP('localhost')
    content = """To: jcaesar@example.org From: soothsayer@example.org Beware the Ides of March. """
    server.sendmail('soothsayer@example.org', 'jcaesar@example.org', content)
    server.quit()

def test23():
    a = [1, 4, 7]
    b = [2, 5, 8]
    c = [3, 6, 9, 10]
    a = list(zip(a, b, c)); p(a)

    list = [
        [1, 4, 7],
        [2, 5, 8],
        [3, 6, 9, 10]
    ]

    a = [[row[i] for row in list] for i in range(3)]; p(a)

def test24():

    # 随机数
    import random
    a = random.choice(['apple', 'pear', 'banana']); p(a)
    a = random.sample(range(100), 10); p(a)
    a = random.random(); p(a)
    a = random.randrange(6); p(a)

def test25():
    # 互联网访问
    from urllib.request import urlopen
    with urlopen('http://tycho.usno.navy.mil/cgi-bin/timer.pl') as response:
        for line in response:
            line = line.decode('utf-8')
            if 'EST' in line or 'EDT' in line:
                print(line)

def test26():
    # 正则匹配
    import re
    re.findall(r'\bf[a-z]*', 'which foot or hand fell fastest')
    re.sub(r'(\b[a-z]+) \1', r'\1', 'cat in the the hat')


def test27():
    class Reverse:

        """迭代一个并反转一个字符串"""

        def __init__(self, string):
            self.string = string
            self.index = len(string)

        # 可以直接返回一个带有 __next__() 方法的对象
        def __iter__(self):
            return self

        def __next__(self):
            # 当index 为0时，遍历完成
            if self.index == 0: raise StopIteration

            # 返回最后的索引
            self.index = self.index - 1
            return self.string[self.index]

    for char in iter(Reverse('abc')): print(char)

    def reverse(data):
        for index in range(len(data) - 1, -1, -1):
            yield data[index]

    for char in reverse('abc'): print(char)


def test28():

    # 推导到达式
    a = list(map(lambda x: x ** 2, range(10)))
    a = [(x, y) for x in [1, 2, 3] for y in [3, 1, 4] if x != y]; print(a)
    a = [(x, x ** 2) for x in range(6)]; print(a)
    a = [x * x for x in range(1, 11) if x % 2 == 0]; print(a)
    a = [m + n for m in 'ABC' for n in 'XYZ']; print(a)
    a = [str(k) + v for k, v in d.items()]; print(a)



def test29():

    a = ['a', 'c', 'b', 'e', 'd']
    a = sorted(a, reverse=True); p(a)
    a = sorted(a, key=len, reverse=True); p(a) # 按照长度排序，当长度一样时，不会进行位置调换

    import bisect
    import sys
    haystack    = [1, 4, 5, 6, 8, 12, 15, 20, 21, 23, 23, 26, 29, 30]
    needles     = [0, 1, 2, 5, 8, 10, 22, 23, 29, 30, 31]

    rowfmt = '{0:2d} @ {1:2d}    {2}{0:<2d}'
    def demo(bisect_fn):
        for needle in reversed(needles):
            position = bisect_fn(haystack, needle)
            offset = position * '  |'
            print(rowfmt.format(needle, position, offset))

    bisect_fn = bisect.bisect
    #bisect_fn = bisect.bisect_left;

    print('haystack ->', ' '.join('%2d' % n for n in haystack))
    demo(bisect_fn)




import collections
isinstance('abc', collections.Iterable) #是否可迭代

# enumerate 把列表变成 索引与元素的键值对儿
for i, val in enumerate(['a', 'b', 'c']):
    print(i, val)

exit()
class B(Exception):
    pass

class C(B):
    pass

class D(C):
    pass

import sys

try:

    f = open('myfile.txt')
    s = f.readline()
    i = int(s.strip())

except OSError as err:
    print("OS error: {0}".format(err))

except ValueError:
    print("Could not convert data to an integer.")

except:
    print("Unexpected error:", sys.exc_info()[0])
    raise



for arg in sys.argv[1:]:
    try:
        f = open(arg, 'r')
    except OSError:
        print('cannot open', arg)
    else:
        #success
        print(arg, 'has', len(f.readlines()), 'lines')
        f.close()


try:
    raise Exception('spam', 'eggs')
except Exception as e:
    print(e)

def divide(x, y):
    try:
        result = x / y
    except ZeroDivisionError:
        print("division by zero!")
    else:
        print("result is", result)
    finally:
        print("executing finally clause")



table = {'Sjoerd': 4127, 'Jack': 4098, 'Dcab': 8637678}
print('Jack: {0[Jack]:d}; Sjoerd: {0[Sjoerd]:d}; Dcab: {0[Dcab]:d}'.format(table))


import argparse
from getpass import getuser
parser = argparse.ArgumentParser(description='An argparse example.')
parser.add_argument('name', nargs='?', default=getuser(), help='The name of someone to greet.')
parser.add_argument('--verbose', '-v', action='count')
args = parser.parse_args()
greeting = ["Hi", "Hello", "Greetings! its very nice to meet you"][args.verbose % 3]
print(f'{greeting}, {args.name}')
if not args.verbose:
    print('Try running this again with multiple "-v" flags!')



关键字参数:函数调用时,通过指定参数名称的方式传递参数.函数定义时,可以定义一个 **keywords 字典来接收说没有名称的关键字参数
解包参数列表: 将任意长度参数列表参数传入函数,会自动解包,填补到相应参数位置.
args = [3, 6]
list(range(*args))


也可以使用形如 kwarg=value 的 关键字参数 来调用函数。

函数标注: 存放在函数的 __annotations__ 属性中.  形参标注格式: args : type [=value]  返回值标注 -> type  完整:   (args : type [=value] ) -> type:

函数默认值只会执行一次。 在默认值为可变对象时很重要, 会和预期表现不同。

创建空集合只能用 set() 因为 {} 是创建一个空字典

__all__ = ["echo", "surround", "reverse"]  表示 from sound.effects import * 将导入 sound 包的三个命名子模块

异常先从基类抛出.

抛出异常
raise NameError('HiThere')

如果传递的是一个异常类，它将通过调用没有参数的构造函数来隐式实例化:
raise ValueError

import mymodule * 只会导入公共成员


mymodule.py
def sayhi():
    print('hello')
__version__ = '1.0'

imoport mymodule
mymodule.sayhi()
mymodule.__version__


# 弱引用
import weakref, gc

class A:

    def __init__(self, value):
        self.value = value

    def __repr__(self):
        return str(self.value)

a = A(10)

d = weakref.WeakValueDictionary()
d['primary'] = a

del a
gc.collect() # 垃圾回收
print(d) # 这个也不存在了



import collections
from random import choice

Card = collections.namedtuple('Card', ['rank', 'suit'])


class FrenchDeck:

    ranks = [str(n) for n in range(2, 11)] + list('JQKA')
    suits = 'spades diamonds clubs hearts'.split()

    def __init__(self):
        self._card = [Card(rank, suit) for suit in self.suits for rank in self.ranks]

    def __len__(self):
        return len(self._card)

    def __getitem__(self, position):
        return self._card[position]


suit_values = dict(spades=3, diamonds=1, clubs=0, hearts=2)


def spades_high(card):
    rank_value = FrenchDeck.ranks.index(card.rank)
    return rank_value * len(suit_values) + suit_values[card.suit]



deck = FrenchDeck()

#print(deck[-1])
#print(choice(deck))
#print(deck[12::13])

for card in sorted(deck, key=spades_high):
    print(card)


########
from math import hypot

class Vector:

    def __init__(self, x=0, y=0):
        self.x = x
        self.y = y

    def __repr__(self):
        return 'Vector(%r, %r)' % (self.x, self.y)

    def __abs__(self):
        return hypot(self.x, self.y)

    def __bool__(self):
        return bool(abs(self))

    def __add__(self, other):
        x = self.x + other.x
        y = self.y + other.y
        return Vector(x, y)

    def __mul__(self, scalar):
        return Vector(self.x * scalar, self.y * scalar)

v1 = Vector(2, 4)
v2 = Vector(2, 1)
#print(v1 + v2)

v = Vector(3, 4)
#print(abs(v))
#print(v * 3)
#print( abs(v * 3) )


symbols = '$%^&*'
beyond_ascii = [ord(s) for s in symbols if ord(s) > 3]
#print(beyond_ascii)

beyond_ascii = list(filter(lambda c : c > 3, map(ord, symbols)))

colors = ['black', 'white']
sizes = ['s', 'm', 'l']

tshirts = [(color, size) for color in colors for size in sizes]
print(tshirts)

symbols = '$%^&*'

a = tuple(ord(symbol) for symbol in symbols)
p(a)

import array
b = array.array("i", (ord(symble) for symble in symbols))
#p(b)

for tshirt in ('%s %s' % (c, s) for c in colors for s in sizes):
    print(tshirt)

city, year = ('tokey', 2003)

travelers = [('USA', '3301'), ('BRA', '3522')]
for passport in sorted(travelers):
    print('%s/%s' % passport)

for country, _ in travelers:
    print(country)

t = (40, 3)
p(divmod(*t))

a, *more, b = range(5)
p(more)

metro_areas = [
    ('Tokyo', 'JP', 36, (35, 139)),
    ('Delhi NCR', 'IN', 21, (28, 77)),
    ('Mexico City', 'MX', 20, (19, -99)),
    ('New York-Newark', 'US', 28, (40, -74))
]

print( '{:15} | {:^9} | {:^9}'.format('city', 'lat.', 'long.') )
fmt = '{:15} | {:9.4f} | {:9.4f}'
for name, cc, pop, (lat, long) in metro_areas:
    if long <= 0:
        print(fmt.format(name, lat, long))


# namedtuple 创建一个和tuple类似的对象，而且对象拥有可访问的属性
import collections
#创建 City 类，包属性有国家/城市名/坐标rint
City = collections.namedtuple('City', 'country name coordinate')
city = City('中国', '武汉',  (35, 139))

p(City._fields)
p(city.country, city[1], city.coordinate)


Latlong = collections.namedtuple('LatLong', 'lat long')
delhi = City._make(('中国', '深圳', Latlong(28, 77)))
p(delhi._asdict())

for key, value in delhi._asdict().items():
    print(key, ':', value)




#!/usr/bin/env python
import random
import re
import requests
from bs4 import BeautifulSoup


def log(*args):

    file = "debug.log"
    date = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())

    with open(file, 'a+', encoding='utf-8') as fp:
        count = len(args)
        for content in args:
            if not isinstance(content, str): content = str(content)
            fp.write(content + '\n-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------\n')

        fp.write('==============================================================================' + date + '==============================================================================\n')
        fp.write('-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------\n')

    print('写入日志：'+ file +'\n')

def stop(one, *args):
    if(args):
        print(one, args)
    else:
        print(one)
    exit()


from PIL import Image
from io import BytesIO

i = Image.open(BytesIO(r.content))


r = requests.get('https://api.github.com/events')
r.json()


r = requests.get('https://api.github.com/events', stream=True)
r.raw
r.raw.read(10)

with open(filename, 'wb') as fd:
    for chunk in r.iter_content(chunk_size):
        fd.write(chunk)


import json
url = 'https://api.github.com/some/endpoint'
payload = {'some': 'data'}
r = requests.post(url, data=json.dumps(payload))
r = requests.post(url, json=payload)



上传文件

url = 'http://httpbin.org/post'
files = {'file': open('report.xls', 'rb')}
r = requests.post(url, files=files)

#显式地设置文件名，文件类型和请求头
url = 'http://httpbin.org/post'
files = {'file': ('report.xls', open('report.xls', 'rb'), 'application/vnd.ms-excel', {'Expires': '0'})}
r = requests.post(url, files=files)

#通过文件方式来发送字符串
url = 'http://httpbin.org/post'
files = {'file': ('report.csv', 'some,data,to,send\nanother,row,to,send\n')}
r = requests.post(url, files=files)


#响应状态码
r.status_code == requests.codes.ok

# Response.raise_for_status() 通过抛出异常
r = requests.get('http://httpbin.org/status/404')
r.raise_for_status()

#Cookie
cookies = dict(cookies_are='working')
r = requests.get(url, cookies=cookies)
r.cookies['example_cookie_name']

# cookies.RequestsCookieJar 和字典类似，但接口更为完整，适合跨域名跨路径使用
jar = requests.cookies.RequestsCookieJar()
jar.set('tasty_cookie', 'yum', domain='httpbin.org', path='/cookies')
jar.set('gross_cookie', 'blech', domain='httpbin.org', path='/elsewhere')
url = 'http://httpbin.org/cookies'
r = requests.get(url, cookies=jar)
r.text
r.history

#超时,是指超时没有应答
requests.get('http://github.com', timeout=0.001)


#会话对象让你能够跨请求保持某些参数。
s = requests.Session()
s.get('http://httpbin.org/cookies/set/sessioncookie/123456789')
r = s.get("http://httpbin.org/cookies")
print(r.text)


#会话也可用来为请求方法提供缺省数据
s = requests.Session()
s.auth = ('user', 'pass')
s.headers.update({'x-test': 'true'})

# both 'x-test' and 'x-test222' are sent
s.get('http://httpbin.org/headers', headers={'x-test222': 'true'})


#cookie  不会被跨请求保持
s = requests.Session()

r = s.get('http://httpbin.org/cookies', cookies={'from-my': 'browser'})
print(r.text)
# '{"cookies": {"from-my": "browser"}}'
r = s.get('http://httpbin.org/cookies')
print(r.text)
# '{"cookies": {}}'



#python -m pip install SomePackage
#python -m pip install "SomePackage>=1.0.4"  # minimum version
#python -m pip install --upgrade SomePackage


#!/usr/bin/env python3
# -*- coding: cp1252 -*-

# ** 乘方

#在交互模式下，上一次打印出来的表达式被赋值给变量 _

#不希望前置了 \ 的字符转义成特殊字符，可以使用 原始字符串 方式，在引号前添加 r 即可
print(r'C:\some\name')

#跨行连续输入是用三重引号："""...""" 或 '''...''' 行尾 \ 表示忽略换行
print("""\
Usage: thingy [OPTIONS]
     -h   Display this usage message
""")

#字符串重复
3 * 'un'

#字符串切片：左闭右开区间。
str(2,5)

#切片中的越界索引会被自动处理


通过组合一些值得到多种 复合 数据类型

列表
squares = [1, 4, 9, 16, 25]

list.append(x) 追加元素。相当于 a[len(a):] = [x]
list.extend(iterable) 使用可迭代对象中的所有元素来扩展列表。相当于 a[len(a):] = iterable
list.insert(i, x)在给定的位置插入一个元素。第一个参数是要插入的元素的索引，所以 a.insert(0, x) 插入列表头部， a.insert(len(a), x)
list.remove(x) 移除列表中第一个值为 x 的元素
list.pop([i]) 删除列表中给定位置的元素并返回它。如果没有给定位置，a.pop() 将会删除并返回列表中的最后一个元素。（ 方法签名中 i 两边的方括号表示这个参数是可选的，而不是要你输入方括号。你会在 Python 参考库中经常看到这种表示方法)
list.clear() 移除列表中的所有元素。等价于 del a[:]
list.index(x[, start[, end]])返回列表中第一个值为 x 的元素的从零开始的索引。如果没有这样的元素将会抛出 ValueError 异常。可选参数 start 和 end 是切片符号，用于将搜索限制为列表的特定子序列。返回的索引是相对于整个序列的开始计算的，而不是 start 参数
list.count(x) 返回元素 x 在列表中出现的次数
list.sort(key=None, reverse=False) 对列表中的元素进行排序（参数可用于自定义排序，解释请参见 sorted()）
list.reverse() 翻转列表中的元素
list.copy() 返回列表的一个浅拷贝，等价于 a[:]

注意：并非所有数据或可以排序或比较。 例如： [None, 'hello', 10] 因为整数不能与字符串比较，而 None 不能与其他类型比较。

列表也可以用作队列，其中先添加的元素被最先取出 (“先进先出”)，在列表的末尾添加和弹出元素非常快，但是在列表的开头插入或弹出元素却很慢(因为所有的其他元素都必须移动一位)。

from collections import deque
queue = deque(["Eric", "John", "Michael"])
queue.append("Terry")           # Terry arrives
queue.append("Graham")          # Graham arrives
queue.popleft()                 # The first to arrive now leaves


列表推导式
提供了一个更简单的创建列表的方法。常见的用法是把某种操作应用于序列或可迭代对象的每个元素上，然后使用其结果来创建列表，或者通过满足某些特定条件元素来创建子序列。

squares = list(map(lambda x: x**2, range(10)))
squares = [x**2 for x in range(10)]
[(x, y) for x in [1,2,3] for y in [3,1,4] if x != y]

#列表推导式将交换其行和列
matrix = [[1, 2, 3, 4], [5, 6, 7, 8], [9, 10, 11, 12],]
[[row[i] for row in matrix] for i in range(4)]
list(zip(*matrix))

del list[2:4]

元组
由几个被逗号隔开的值组成,可以带括号。
tuple = 12345, 54321, 'hello!'
tuple = ('a', 'b', 'c')
tuple = ()
tuple = ('a')
print(tuple)

语句 tuple = 12345, 54321, 'hello!' 是元组打包，值被打包进元组。其逆操作叫序列解包 x, y, z = tuple

集合
花括号或 set() 函数来创建集合。注意：要创建空集合只能用 set() 因为一对花括号是创建空字典的。
sets = {'apple', 'orange', 'apple', 'pear', 'orange', 'banana'}
sets = set('apple,orange, apple, pear, orange')
print(sets)

set_a - set_b   # 补集
set_a | set_b   # 并集
set_a & set_b   # 交集
set_a ^ set_b   # "异集"

a = {x for x in 'abracadabra' if x not in 'abc'}

字典
在其他语言里可能会被叫做联合数组。与以连续整数为索引的序列不同，字典是以关键字为索引，关键字可以是任意不可变类型，通常是字符串或数字。如果元组包含了可变对象，那它就不能用作关键字。

tel = {'jack': 4098, 'sape': 4139}
tel['guido'] = 4127
dict([('sape', 4139), ('guido', 4127), ('jack', 4098)])
dict(sape=4139, guido=4127, jack=4098)
{x: x**2 for x in (2, 4, 6)}


list(d) 返回字典中所有的键名
sorted(d) 按键名排序
'a' in { 'a':123 } 检查字典中是否存在一个特定键

#遍历字典元素
knights = {'gallahad': 'the pure', 'robin': 'the brave'}
for k, v in knights.items():
    print(k, v)

# enumerate() 函数可以从列表中将索引位置和值同时取出
for i, v in enumerate(['tic', 'tac', 'toe']):
    print(i, v)

#用 zip() 函数将其内元素一一匹配
questions = ['name', 'quest', 'favorite color']
answers = ['lancelot', 'the holy grail', 'blue']

for q, a in zip(questions, answers):
    print('What is your {0}?  It is {1}.'.format(q, a))

import math
raw_data = [56.2, float('NaN'), 51.7, 55.3, 52.5, float('NaN'), 47.8]
filtered_data = []
for value in raw_data:
    if not math.isnan(value):
filtered_data.append(value)

from fibo import * 这会调入所有非以下划线（_）开头的名称。
import fibo as fb
from fibo import fib as fibonacci

#出于效率的考虑，每个模块在每个解释器会话中只被导入一次。你可以用  importlib.reload()
import importlib;
importlib.reload(modulename)

模块搜索过程：先找内置模块，然后在 sys.path 中查找。

__pycache__ 目录里缓存了每个模块的编译后版本
compileall 模块可以为一个目录下的所有模块创建.pyc文件。

import sys
sys.ps1
sys.path.append('/ufs/guido/lib/python')

__init__.py 可以只是一个空文件，但它也可以执行包的初始化代码。如__all__ = ["echo", "surround", "reverse"] 这意味着 from sound.effects import * 将导入 sound 包的三个命名子模块

子包通过相对路径相互导入：
例如，从 surround 模块，你可以使用:
from . import echo
from .. import formats
from ..filters import equalizer


格式化
f'Results of the {year} {event}'

yes_votes = 42_572_654
no_votes = 43_132_495
percentage = yes_votes / (yes_votes + no_votes)
'{:-9} YES votes  {:2.2%}'.format(yes_votes, percentage)

import math
print(f'The value of pi is approximately {math.pi:.3f}.')


table = {'Sjoerd': 4127, 'Jack': 4098, 'Dcab': 7678}
for name, phone in table.items():
    print(f'{name:10} ==> {phone:10d}')


其他的修饰符可用于在格式化之前转化值。 '!a' 应用 ascii() ，'!s' 应用 str()，还有 '!r' 应用 repr():
animals = 'eels'
print(f'My hovercraft is full of {animals}.')
print(f'My hovercraft is full of {animals!r}.')

str.format() 方法的基本用法如下所示:
print('We are the {} who say "{}!"'.format('knights', 'Ni'))
print('{1} and {0}'.format('spam', 'eggs'))
print('This {food} is {adjective}.'.format(food='spam', adjective='absolutely horrible'))
#位置和关键字参数可以任意组合:
print('The story of {0}, {1}, and {other}.'.format('Bill', 'Manfred', other='Georg'))

table = {'Sjoerd': 4127, 'Jack': 4098, 'Dcab': 8637678}
print('Jack: {0[Jack]:d}; Sjoerd: {0[Sjoerd]:d}; ' 'Dcab: {0[Dcab]:d}'.format(table))

这也可以通过使用 '**' 符号将 table 作为关键字参数传递。
table = {'Sjoerd': 4127, 'Jack': 4098, 'Dcab': 8637678}
print('Jack: {Jack:d}; Sjoerd: {Sjoerd:d}; Dcab: {Dcab:d}'.format(**table))

for x in range(1, 11):
    print('{0:2d} {1:3d} {2:4d}'.format(x, x*x, x*x*x))


import json
json.dumps([1, 'simple', 'list'])

try:
    raise Exception('spam', 'eggs')
except Exception as inst:
    print(type(inst))    # the exception instance
    print(inst.args)     # arguments stored in .args
    print(inst)          # __str__ allows args to be printed directly,
    # but may be overridden in exception subclasses
    x, y = inst.args     # unpack args
    print('x =', x)
    print('y =', y)

try:
    result = x / y
except ZeroDivisionError:
    print("division by zero!")
else:
    print("result is", result)
finally:
    print("executing finally clause")



with open("myfile.txt") as f:
    for line in f:
        print(line, end="")


import builtins
dir(builtins)

for w in words:
    print(w, len(w))

# Strategy:  Iterate over a copy
for user, status in users.copy().items():
    if status == 'inactive':
        del users[user]


# Strategy:  Create a new collection
active_users = {}
for user, status in users.items():
    if status == 'active':
        active_users[user] = status


for i in range(5):
    print(i)

a = ['Mary', 'had', 'a', 'little', 'lamb']
for i in range(len(a)):
    print(i, a[i])




def fib(n):    # write Fibonacci series up to n
    """Print a Fibonacci series up to n."""
    a, b = 0, 1
    while a < n:
        print(a, end=' ')
        a, b = b, a+b
    print()


i = 5
def f(arg=i):
    print(arg)

i = 6
f()


def parrot(voltage, state='a stiff', action='voom', type='Norwegian Blue'):
    print(voltage, state, action, type)

parrot(1, 2, 3)

/ 仅限位置参数
* 仅限关键字参数

def pos_only_arg(arg, /):
    print(arg)

def kwd_only_arg(*, arg):
    print(arg)

def combined_example(pos_only, /, pos_or_kwd, *, kwd_only):
    pass

仅限位置形参的名称可以在 **kwds 中使用而不产生歧义。
def foo(name, /, **kwds):
    return 'name' in kwds


解包参数列表(*|**)
args = [3, 6]
list(range(*args))

d = {"voltage": "four million", "state": "bleedin' demised", "action": "VOOM"}
parrot(**d)


Lambda 表达式

def make_incrementor(n):
    return lambda x: x + n
f = make_incrementor(42)
f(1) #43


pairs = [(1, 'one'), (2, 'two'), (3, 'three'), (4, 'four')]
pairs.sort(key=lambda pair: pair[1])
pairs  #[(4, 'four'), (1, 'one'), (3, 'three'), (2, 'two')]


文档字符串
def my_function():
    """Do nothing, but document
    No, really, it doesn't do anything.
    """
    pass

print(my_function.__doc__)


函数标注
函数标注 是关于用户自定义函数中使用的类型的完全可选元数据信息

def f(ham: str, eggs: str = 'eggs') -> str:
    print("Annotations:", f.__annotations__)

f('spam')



class Myclass:
    """一个简单的示例"""
    i = 124234

    def __init__(self):
        super().__init__()
        self.i = "afsfsfs"

    def f(self):
        print(self.i)

x = Myclass()
x.f()


实例变量用于每个实例各自拥有，而类变量用于所有实例共享

class Dog:

    tricks = []             # mistaken use of a class variable

    def __init__(self, name):
        self.name = name

    def add_trick(self, trick):
        self.tricks.append(trick)

d = Dog('Fido')
e = Dog('Buddy')
d.add_trick('roll over')
e.add_trick('play dead')

print(e.name, e.tricks)

it = iter('abcdefg')

print(next(it), next(it), next(it), next(it), sep="--")


class Reverse:
    """Iterator for looping over a sequence backwards."""
    def __init__(self, data):
        self.data = data
        self.index = len(data)

    def __iter__(self):
        return self

    def __next__(self):

        if self.index == 0:
            raise StopIteration

        self.index = self.index - 1

        return self.data[self.index]

rev = Reverse('spam')

iter(rev)

for char in rev:
    print(char)



sum((i for i in range(4)))  # 0 + 1 + 2 + 3  = 6

data = 'golf'
list(data[i] for i in range(len(data)-1, -1, -1))

import os
os.getcwd()
import shutil
shutil.copyfile('data.db', 'archive.db')
shutil.move('/build/executables', 'installdir')

import glob
glob.glob('*.py')

import sys
print(sys.argv)

argparse 模块提供了一种更复杂的机制来处理命令行参数。 以下脚本可提取一个或多个文件名，并可选择要显示的行数:

import argparse

parser = argparse.ArgumentParser(prog = 'top',
                                 description = 'Show top lines from each file')
parser.add_argument('filenames', nargs='+')
parser.add_argument('-l', '--lines', type=int, default=10)
args = parser.parse_args()
print(args)
当在通过 python top.py --lines=5 alpha.txt beta.txt 在命令行运行时，该脚本会将 args.lines 设为 5 并将 args.filenames 设为 ['alpha.txt', 'beta.txt']。

sys 模块还具有 stdin ， stdout 和 stderr 的属性。后者对于发出警告和错误消息非常有用，即使在 stdout 被重定向后也可以看到它们:
sys.stderr.write('Warning, log file not found starting a new one\n')

终止脚本的最直接方法是使用 sys.exit()

re 模块为高级字符串处理提供正则表达式工具
import re
re.findall(r'\bf[a-z]*', 'which foot or hand fell fastest')
re.sub(r'(\b[a-z]+) \1', r'\1', 'cat in the the hat')

当只需要简单的功能时，首选字符串方法因为它们更容易阅读和调试:
'tea for too'.replace('too', 'two')


statistics 模块计算数值数据的基本统计属性（均值，中位数，方差等）
SciPy项目 <https://scipy.org> 有许多其他模块用于数值计算


互联网访问
有许多模块可用于访问互联网和处理互联网协议。其中两个最简单的 urllib.request 用于从URL检索数据，以及 smtplib 用于发送邮件:

from urllib.request import urlopen
with urlopen('http://tycho.usno.navy.mil/cgi-bin/timer.pl') as response:
    for line in response:
        line = line.decode('utf-8')  # Decoding the binary data to text.
        if 'EST' in line or 'EDT' in line:  # look for Eastern Time
            print(line)



import smtplib
server = smtplib.SMTP('localhost')
server.sendmail('soothsayer@example.org', 'jcaesar@example.org',
                """To: jcaesar@example.org
                From: soothsayer@example.org
                
                Beware the Ides of March.
                """)
server.quit()



from datetime import date
now = date.today()

数据存档和压缩格式由模块直接支持，包括：zlib, gzip, bz2, lzma, zipfile 和 tarfile
import zlib
s = b'witch which has which witches wrist watch'
len(s)
t = zlib.compress(s)


性能测量
一些Python用户对了解同一问题的不同方法的相对性能产生了浓厚的兴趣。 Python提供了一种可以立即回答这些问题的测量工具。
from timeit import Timer
Timer('t=a; a=b; b=t', 'a=1; b=2').timeit()
Timer('a,b = b,a', 'a=1; b=2').timeit()


质量控制
def average(values):
    """Computes the arithmetic mean of a list of numbers.

    >>> print(average([20, 30, 70]))
    40.0
    """
    return sum(values) / len(values)

import doctest
doctest.testmod()


unittest 模块不像 doctest 模块那样易于使用，但它允许在一个单独的文件中维护更全面的测试集:

import unittest

class TestStatisticalFunctions(unittest.TestCase):

    def test_average(self):
        self.assertEqual(average([20, 30, 70]), 40.0)
        self.assertEqual(round(average([1, 5, 7]), 1), 4.3)
        with self.assertRaises(ZeroDivisionError):
            average([])
        with self.assertRaises(TypeError):
            average(20, 30, 70)

unittest.main()  # Calling from the command line invokes all tests


import bisect
scores = [(100, 'perl'), (200, 'tcl'), (400, 'lua'), (500, 'python')]
bisect.insort(scores, (300, 'ruby'))
print(scores)


from array import array
a = array('H', [1, 10, 7, 2])
print(sum(a))
print(a[1:3])



from docx import Document
from docx.shared import Cm

doc = Document('d:\\input.docx')
paragraph = doc.paragraphs[1]
paragraph.runs
paragraph2 = doc.add_paragraph()
paragraph2.add_run("粗体内容").bold = True

doc.add_page_break()
doc.add_picture('d:\\a.jpg', width=Cm(5))

record = [
    ['学号', '姓名'],
    [101, '张三']
    [102, '李四']
]

table = doc.add_table(rows=3, cols=2)
for row in range(3):
    cells = table.rows[row].cells
    for col in range(2):
        cells[col].text = str(record[row][col])




from bs4 import BeautifulSoup
import requests

html = """
    <span></span>
"""


all = soup.find_all(True)
for tag in all:
    print(tag.name)


soup.p['class']
soup.p.get('class')
soup.p['class']="newClass"

type(soup.a.string)

soup.head.contents[0]

for child in  soup.body.children:
    print child

#所有子节点
for child in soup.descendants:
    print child


for string in soup.strings:
    print(repr(string))

#去除多余空白
for string in soup.stripped_strings:
    print(repr(string))


#所有你节点
content = soup.head.title.string
for parent in  content.parents:
    print parent.name

前后兄弟节点
.next_sibling .next_siblings .previous_sibling .previous_siblings

前后节点，不分层次
.next_element .previous_element .next_elements .previous_elements

find_all( name , attrs , recursive , string , **kwargs )
soup.find_all(["a", "b"])

#自定义过滤器
def has_class_but_no_id(tag):
    return tag.has_attr('class') and not tag.has_attr('id')
soup.find_all(has_class_but_no_id)

soup.find_all(id='link2')

#匹配带 .com
soup.find_all(href=re.compile(".com"))

soup.find_all("a", class_="sister")
data_soup.find_all(data-foo="value")

data_soup.find_all(attrs={"data-foo": "value"})

#innerHtml
soup.find_all(text="Elsie")
soup.find_all("a", limit=2， recursive=False)

soup.get_text()
soup.prettify()


# num = re.sub(r'<div class="pen-l">.*>\s+返回列表\s+</a>', '<div class="pen-l">', content.prettify(), flags=re.S)
# print(num)


from bs4 import BeautifulSoup

soup = BeautifulSoup(open("index.html"))
tag = soup.tag
tag.name
tag.string
tag.string.replace_with("No longer bold")
tag['class']
tag.prettify()

soup.find_all('a')
tag.contents
tag.childre
tag.descendants #对所有tag的子孙节点进行递归循环
tag.strings
tag.stripped_strings
tag.parent
tag.parents
tag.next_sibling
tag.previous_sibling
tag.next_siblings
tag.previous_siblings
tag.next_element
tag.previous_element
tag.next_elements
tag.previous_elements
soup.find()
import re
for tag in soup.find_all(re.compile("^b")):
    print(tag.name)

soup.find_all(["a", "b"])
soup.find_all(True) #过滤字符串节点

def has_class_but_no_id(tag):
    return tag.has_attr('class') and not tag.has_attr('id')

soup.find_all(has_class_but_no_id)

soup.find_all(id='link2')
soup.find_all(href=re.compile("elsie"))
soup.find_all(id=True) #包含id属性
