#!/usr/bin/env python
import random
import re
import requests
from bs4 import BeautifulSoup


def log(*args):

    file = "debug.log"
    date = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())

    with open(file, 'a+', encoding='utf-8') as fp:
        count = len(args)
        for content in args:
            if not isinstance(content, str): content = str(content)
            fp.write(content + '\n-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------\n')

        fp.write('==============================================================================' + date + '==============================================================================\n')
        fp.write('-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------\n')

    print('写入日志：'+ file +'\n')

def stop(one, *args):
    if(args):
        print(one, args)
    else:
        print(one)
    exit()


from PIL import Image
from io import BytesIO

i = Image.open(BytesIO(r.content))


r = requests.get('https://api.github.com/events')
r.json()


r = requests.get('https://api.github.com/events', stream=True)
r.raw
r.raw.read(10)

with open(filename, 'wb') as fd:
    for chunk in r.iter_content(chunk_size):
        fd.write(chunk)


import json
url = 'https://api.github.com/some/endpoint'
payload = {'some': 'data'}
r = requests.post(url, data=json.dumps(payload))
r = requests.post(url, json=payload)



上传文件

url = 'http://httpbin.org/post'
files = {'file': open('report.xls', 'rb')}
r = requests.post(url, files=files)

#显式地设置文件名，文件类型和请求头
url = 'http://httpbin.org/post'
files = {'file': ('report.xls', open('report.xls', 'rb'), 'application/vnd.ms-excel', {'Expires': '0'})}
r = requests.post(url, files=files)

#通过文件方式来发送字符串
url = 'http://httpbin.org/post'
files = {'file': ('report.csv', 'some,data,to,send\nanother,row,to,send\n')}
r = requests.post(url, files=files)


#响应状态码
r.status_code == requests.codes.ok

# Response.raise_for_status() 通过抛出异常
r = requests.get('http://httpbin.org/status/404')
r.raise_for_status()

#Cookie
cookies = dict(cookies_are='working')
r = requests.get(url, cookies=cookies)
r.cookies['example_cookie_name']

# cookies.RequestsCookieJar 和字典类似，但接口更为完整，适合跨域名跨路径使用
jar = requests.cookies.RequestsCookieJar()
jar.set('tasty_cookie', 'yum', domain='httpbin.org', path='/cookies')
jar.set('gross_cookie', 'blech', domain='httpbin.org', path='/elsewhere')
url = 'http://httpbin.org/cookies'
r = requests.get(url, cookies=jar)
r.text
r.history

#超时,是指超时没有应答
requests.get('http://github.com', timeout=0.001)


#会话对象让你能够跨请求保持某些参数。
s = requests.Session()
s.get('http://httpbin.org/cookies/set/sessioncookie/123456789')
r = s.get("http://httpbin.org/cookies")
print(r.text)


#会话也可用来为请求方法提供缺省数据
s = requests.Session()
s.auth = ('user', 'pass')
s.headers.update({'x-test': 'true'})

# both 'x-test' and 'x-test222' are sent
s.get('http://httpbin.org/headers', headers={'x-test222': 'true'})


#cookie  不会被跨请求保持
s = requests.Session()

r = s.get('http://httpbin.org/cookies', cookies={'from-my': 'browser'})
print(r.text)
# '{"cookies": {"from-my": "browser"}}'
r = s.get('http://httpbin.org/cookies')
print(r.text)
# '{"cookies": {}}'



#python -m pip install SomePackage
#python -m pip install "SomePackage>=1.0.4"  # minimum version
#python -m pip install --upgrade SomePackage


#!/usr/bin/env python3
# -*- coding: cp1252 -*-

# ** 乘方

#在交互模式下，上一次打印出来的表达式被赋值给变量 _

#不希望前置了 \ 的字符转义成特殊字符，可以使用 原始字符串 方式，在引号前添加 r 即可
print(r'C:\some\name')

#跨行连续输入是用三重引号："""...""" 或 '''...''' 行尾 \ 表示忽略换行
print("""\
Usage: thingy [OPTIONS]
     -h   Display this usage message
""")

#字符串重复
3 * 'un'

#字符串切片：左闭右开区间。
str(2,5)

#切片中的越界索引会被自动处理


通过组合一些值得到多种 复合 数据类型

列表
squares = [1, 4, 9, 16, 25]

list.append(x) 追加元素。相当于 a[len(a):] = [x]
list.extend(iterable) 使用可迭代对象中的所有元素来扩展列表。相当于 a[len(a):] = iterable 
list.insert(i, x)在给定的位置插入一个元素。第一个参数是要插入的元素的索引，所以 a.insert(0, x) 插入列表头部， a.insert(len(a), x) 
list.remove(x) 移除列表中第一个值为 x 的元素
list.pop([i]) 删除列表中给定位置的元素并返回它。如果没有给定位置，a.pop() 将会删除并返回列表中的最后一个元素。（ 方法签名中 i 两边的方括号表示这个参数是可选的，而不是要你输入方括号。你会在 Python 参考库中经常看到这种表示方法)
list.clear() 移除列表中的所有元素。等价于 del a[:]
list.index(x[, start[, end]])返回列表中第一个值为 x 的元素的从零开始的索引。如果没有这样的元素将会抛出 ValueError 异常。可选参数 start 和 end 是切片符号，用于将搜索限制为列表的特定子序列。返回的索引是相对于整个序列的开始计算的，而不是 start 参数
list.count(x) 返回元素 x 在列表中出现的次数
list.sort(key=None, reverse=False) 对列表中的元素进行排序（参数可用于自定义排序，解释请参见 sorted()）
list.reverse() 翻转列表中的元素
list.copy() 返回列表的一个浅拷贝，等价于 a[:]

注意：并非所有数据或可以排序或比较。 例如： [None, 'hello', 10] 因为整数不能与字符串比较，而 None 不能与其他类型比较。

列表也可以用作队列，其中先添加的元素被最先取出 (“先进先出”)，在列表的末尾添加和弹出元素非常快，但是在列表的开头插入或弹出元素却很慢(因为所有的其他元素都必须移动一位)。

from collections import deque
queue = deque(["Eric", "John", "Michael"])
queue.append("Terry")           # Terry arrives
queue.append("Graham")          # Graham arrives
queue.popleft()                 # The first to arrive now leaves


列表推导式
提供了一个更简单的创建列表的方法。常见的用法是把某种操作应用于序列或可迭代对象的每个元素上，然后使用其结果来创建列表，或者通过满足某些特定条件元素来创建子序列。

squares = list(map(lambda x: x**2, range(10)))
squares = [x**2 for x in range(10)]
[(x, y) for x in [1,2,3] for y in [3,1,4] if x != y]

#列表推导式将交换其行和列
matrix = [[1, 2, 3, 4], [5, 6, 7, 8], [9, 10, 11, 12],]
[[row[i] for row in matrix] for i in range(4)]
list(zip(*matrix))

del list[2:4]

元组
由几个被逗号隔开的值组成,可以带括号。
tuple = 12345, 54321, 'hello!'
tuple = ('a', 'b', 'c')
tuple = ()
tuple = ('a')
print(tuple)

语句 tuple = 12345, 54321, 'hello!' 是元组打包，值被打包进元组。其逆操作叫序列解包 x, y, z = tuple

集合
花括号或 set() 函数来创建集合。注意：要创建空集合只能用 set() 因为一对花括号是创建空字典的。
sets = {'apple', 'orange', 'apple', 'pear', 'orange', 'banana'}
sets = set('apple,orange, apple, pear, orange')
print(sets)

set_a - set_b   # 补集
set_a | set_b   # 并集
set_a & set_b   # 交集
set_a ^ set_b   # "异集"

a = {x for x in 'abracadabra' if x not in 'abc'}

字典
在其他语言里可能会被叫做联合数组。与以连续整数为索引的序列不同，字典是以关键字为索引，关键字可以是任意不可变类型，通常是字符串或数字。如果元组包含了可变对象，那它就不能用作关键字。

tel = {'jack': 4098, 'sape': 4139}
tel['guido'] = 4127
dict([('sape', 4139), ('guido', 4127), ('jack', 4098)])
dict(sape=4139, guido=4127, jack=4098)
{x: x**2 for x in (2, 4, 6)}


list(d) 返回字典中所有的键名
sorted(d) 按键名排序
'a' in { 'a':123 } 检查字典中是否存在一个特定键

#遍历字典元素
knights = {'gallahad': 'the pure', 'robin': 'the brave'}
for k, v in knights.items():
    print(k, v)

# enumerate() 函数可以从列表中将索引位置和值同时取出
for i, v in enumerate(['tic', 'tac', 'toe']):
    print(i, v)

#用 zip() 函数将其内元素一一匹配
questions = ['name', 'quest', 'favorite color']
answers = ['lancelot', 'the holy grail', 'blue']

for q, a in zip(questions, answers):
    print('What is your {0}?  It is {1}.'.format(q, a))

import math
raw_data = [56.2, float('NaN'), 51.7, 55.3, 52.5, float('NaN'), 47.8]
filtered_data = []
for value in raw_data:
    if not math.isnan(value):
        filtered_data.append(value)

from fibo import * 这会调入所有非以下划线（_）开头的名称。 
import fibo as fb
from fibo import fib as fibonacci

#出于效率的考虑，每个模块在每个解释器会话中只被导入一次。你可以用  importlib.reload()
import importlib;
importlib.reload(modulename)

模块搜索过程：先找内置模块，然后在 sys.path 中查找。

__pycache__ 目录里缓存了每个模块的编译后版本
compileall 模块可以为一个目录下的所有模块创建.pyc文件。

import sys
sys.ps1
sys.path.append('/ufs/guido/lib/python')

__init__.py 可以只是一个空文件，但它也可以执行包的初始化代码。如__all__ = ["echo", "surround", "reverse"] 这意味着 from sound.effects import * 将导入 sound 包的三个命名子模块

子包通过相对路径相互导入：
例如，从 surround 模块，你可以使用:
from . import echo
from .. import formats
from ..filters import equalizer


格式化
f'Results of the {year} {event}'

yes_votes = 42_572_654
no_votes = 43_132_495
percentage = yes_votes / (yes_votes + no_votes)
'{:-9} YES votes  {:2.2%}'.format(yes_votes, percentage)

import math
print(f'The value of pi is approximately {math.pi:.3f}.')


table = {'Sjoerd': 4127, 'Jack': 4098, 'Dcab': 7678}
for name, phone in table.items():
    print(f'{name:10} ==> {phone:10d}')


其他的修饰符可用于在格式化之前转化值。 '!a' 应用 ascii() ，'!s' 应用 str()，还有 '!r' 应用 repr():
animals = 'eels'
print(f'My hovercraft is full of {animals}.')
print(f'My hovercraft is full of {animals!r}.')

str.format() 方法的基本用法如下所示:
print('We are the {} who say "{}!"'.format('knights', 'Ni'))
print('{1} and {0}'.format('spam', 'eggs'))
print('This {food} is {adjective}.'.format(food='spam', adjective='absolutely horrible'))
#位置和关键字参数可以任意组合:
print('The story of {0}, {1}, and {other}.'.format('Bill', 'Manfred', other='Georg'))

table = {'Sjoerd': 4127, 'Jack': 4098, 'Dcab': 8637678}
print('Jack: {0[Jack]:d}; Sjoerd: {0[Sjoerd]:d}; ' 'Dcab: {0[Dcab]:d}'.format(table))

这也可以通过使用 '**' 符号将 table 作为关键字参数传递。
table = {'Sjoerd': 4127, 'Jack': 4098, 'Dcab': 8637678}
print('Jack: {Jack:d}; Sjoerd: {Sjoerd:d}; Dcab: {Dcab:d}'.format(**table))

for x in range(1, 11):
    print('{0:2d} {1:3d} {2:4d}'.format(x, x*x, x*x*x))


import json
json.dumps([1, 'simple', 'list'])

try:
    raise Exception('spam', 'eggs')
except Exception as inst:
    print(type(inst))    # the exception instance
    print(inst.args)     # arguments stored in .args
    print(inst)          # __str__ allows args to be printed directly,
                         # but may be overridden in exception subclasses
    x, y = inst.args     # unpack args
    print('x =', x)
    print('y =', y)

try:
    result = x / y
except ZeroDivisionError:
    print("division by zero!")
else:
    print("result is", result)
finally:
    print("executing finally clause")



with open("myfile.txt") as f:
    for line in f:
        print(line, end="")


import builtins
dir(builtins)  

for w in words:
    print(w, len(w))

# Strategy:  Iterate over a copy
for user, status in users.copy().items():
    if status == 'inactive':
        del users[user]


# Strategy:  Create a new collection
active_users = {}
for user, status in users.items():
    if status == 'active':
        active_users[user] = status


for i in range(5):
    print(i)

a = ['Mary', 'had', 'a', 'little', 'lamb']
for i in range(len(a)):
    print(i, a[i])




def fib(n):    # write Fibonacci series up to n
    """Print a Fibonacci series up to n."""
    a, b = 0, 1
    while a < n:
        print(a, end=' ')
        a, b = b, a+b
    print()


i = 5
def f(arg=i):
    print(arg)

i = 6
f()


def parrot(voltage, state='a stiff', action='voom', type='Norwegian Blue'):
    print(voltage, state, action, type)

parrot(1, 2, 3)

/ 仅限位置参数
* 仅限关键字参数

def pos_only_arg(arg, /):
     print(arg)

def kwd_only_arg(*, arg):
    print(arg)

def combined_example(pos_only, /, pos_or_kwd, *, kwd_only):
    pass

仅限位置形参的名称可以在 **kwds 中使用而不产生歧义。
def foo(name, /, **kwds):
    return 'name' in kwds


解包参数列表(*|**)
args = [3, 6]
list(range(*args))

d = {"voltage": "four million", "state": "bleedin' demised", "action": "VOOM"}
parrot(**d)


Lambda 表达式

def make_incrementor(n):
    return lambda x: x + n
f = make_incrementor(42)
f(1) #43


pairs = [(1, 'one'), (2, 'two'), (3, 'three'), (4, 'four')]
pairs.sort(key=lambda pair: pair[1])
pairs  #[(4, 'four'), (1, 'one'), (3, 'three'), (2, 'two')]


文档字符串
def my_function():
    """Do nothing, but document
    No, really, it doesn't do anything.
    """
    pass

print(my_function.__doc__)


函数标注
函数标注 是关于用户自定义函数中使用的类型的完全可选元数据信息

def f(ham: str, eggs: str = 'eggs') -> str:
    print("Annotations:", f.__annotations__)

f('spam')



class Myclass:
    """一个简单的示例"""
    i = 124234

    def __init__(self):
        super().__init__()
        self.i = "afsfsfs"

    def f(self):
        print(self.i)

x = Myclass()
x.f()


实例变量用于每个实例各自拥有，而类变量用于所有实例共享

class Dog:

    tricks = []             # mistaken use of a class variable

    def __init__(self, name):
        self.name = name

    def add_trick(self, trick):
        self.tricks.append(trick)

d = Dog('Fido')
e = Dog('Buddy')
d.add_trick('roll over')
e.add_trick('play dead')

print(e.name, e.tricks)

it = iter('abcdefg')

print(next(it), next(it), next(it), next(it), sep="--")


class Reverse:
    """Iterator for looping over a sequence backwards."""
    def __init__(self, data):
        self.data = data
        self.index = len(data)

    def __iter__(self):
        return self

    def __next__(self):

        if self.index == 0:
            raise StopIteration

        self.index = self.index - 1

        return self.data[self.index]

rev = Reverse('spam')

iter(rev)

for char in rev:
    print(char)



sum((i for i in range(4)))  # 0 + 1 + 2 + 3  = 6

data = 'golf'
list(data[i] for i in range(len(data)-1, -1, -1))

import os
os.getcwd()
import shutil
shutil.copyfile('data.db', 'archive.db')
shutil.move('/build/executables', 'installdir')

import glob
glob.glob('*.py')

import sys
print(sys.argv)

argparse 模块提供了一种更复杂的机制来处理命令行参数。 以下脚本可提取一个或多个文件名，并可选择要显示的行数:

import argparse

parser = argparse.ArgumentParser(prog = 'top',
    description = 'Show top lines from each file')
parser.add_argument('filenames', nargs='+')
parser.add_argument('-l', '--lines', type=int, default=10)
args = parser.parse_args()
print(args)
当在通过 python top.py --lines=5 alpha.txt beta.txt 在命令行运行时，该脚本会将 args.lines 设为 5 并将 args.filenames 设为 ['alpha.txt', 'beta.txt']。

sys 模块还具有 stdin ， stdout 和 stderr 的属性。后者对于发出警告和错误消息非常有用，即使在 stdout 被重定向后也可以看到它们:
sys.stderr.write('Warning, log file not found starting a new one\n')

终止脚本的最直接方法是使用 sys.exit() 

re 模块为高级字符串处理提供正则表达式工具
import re
re.findall(r'\bf[a-z]*', 'which foot or hand fell fastest')
re.sub(r'(\b[a-z]+) \1', r'\1', 'cat in the the hat')

当只需要简单的功能时，首选字符串方法因为它们更容易阅读和调试:
'tea for too'.replace('too', 'two')


statistics 模块计算数值数据的基本统计属性（均值，中位数，方差等）
SciPy项目 <https://scipy.org> 有许多其他模块用于数值计算



互联网访问
有许多模块可用于访问互联网和处理互联网协议。其中两个最简单的 urllib.request 用于从URL检索数据，以及 smtplib 用于发送邮件:

from urllib.request import urlopen
with urlopen('http://tycho.usno.navy.mil/cgi-bin/timer.pl') as response:
    for line in response:
        line = line.decode('utf-8')  # Decoding the binary data to text.
        if 'EST' in line or 'EDT' in line:  # look for Eastern Time
            print(line)



import smtplib
server = smtplib.SMTP('localhost')
server.sendmail('soothsayer@example.org', 'jcaesar@example.org',
"""To: jcaesar@example.org
From: soothsayer@example.org

Beware the Ides of March.
""")
server.quit()



from datetime import date
now = date.today()

数据存档和压缩格式由模块直接支持，包括：zlib, gzip, bz2, lzma, zipfile 和 tarfile
import zlib
s = b'witch which has which witches wrist watch'
len(s)
t = zlib.compress(s)


性能测量
一些Python用户对了解同一问题的不同方法的相对性能产生了浓厚的兴趣。 Python提供了一种可以立即回答这些问题的测量工具。
from timeit import Timer
Timer('t=a; a=b; b=t', 'a=1; b=2').timeit()
Timer('a,b = b,a', 'a=1; b=2').timeit()


质量控制
def average(values):
    """Computes the arithmetic mean of a list of numbers.

    >>> print(average([20, 30, 70]))
    40.0
    """
    return sum(values) / len(values)

import doctest
doctest.testmod()


unittest 模块不像 doctest 模块那样易于使用，但它允许在一个单独的文件中维护更全面的测试集:

import unittest

class TestStatisticalFunctions(unittest.TestCase):

    def test_average(self):
        self.assertEqual(average([20, 30, 70]), 40.0)
        self.assertEqual(round(average([1, 5, 7]), 1), 4.3)
        with self.assertRaises(ZeroDivisionError):
            average([])
        with self.assertRaises(TypeError):
            average(20, 30, 70)

unittest.main()  # Calling from the command line invokes all tests


import bisect
scores = [(100, 'perl'), (200, 'tcl'), (400, 'lua'), (500, 'python')]
bisect.insort(scores, (300, 'ruby'))
print(scores)


from array import array
a = array('H', [1, 10, 7, 2])
print(sum(a))
print(a[1:3])



from docx import Document
from docx.shared import Cm

doc = Document('d:\\input.docx')
paragraph = doc.paragraphs[1]
paragraph.runs
paragraph2 = doc.add_paragraph()
paragraph2.add_run("粗体内容").bold = True

doc.add_page_break()
doc.add_picture('d:\\a.jpg', width=Cm(5))

record = [
    ['学号', '姓名'],
    [101, '张三']
    [102, '李四']
]

table = doc.add_table(rows=3, cols=2)
for row in range(3):
    cells = table.rows[row].cells
    for col in range(2):
        cells[col].text = str(record[row][col])


